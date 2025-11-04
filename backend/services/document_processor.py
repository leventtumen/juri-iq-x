import os
import json
import hashlib
from datetime import datetime
from PyPDF2 import PdfReader
import docx
import magic
import nltk
import spacy
from sklearn.feature_extraction.text import TfidfVectorizer
from sklearn.metrics.pairwise import cosine_similarity
import numpy as np
from models.database import db
from models.document import Document, DocumentContent

class DocumentProcessor:
    def __init__(self):
        self.nlp = None
        self._load_nlp_model()
        self._download_nltk_data()
    
    def _load_nlp_model(self):
        try:
            self.nlp = spacy.load('en_core_web_sm')
        except OSError:
            print("Spacy model not found. Please run: python -m spacy download en_core_web_sm")
            self.nlp = None
    
    def _download_nltk_data(self):
        try:
            nltk.data.find('tokenizers/punkt')
        except LookupError:
            nltk.download('punkt')
        
        try:
            nltk.data.find('corpora/stopwords')
        except LookupError:
            nltk.download('stopwords')
    
    def process_document(self, file_path):
        try:
            # Extract text based on file type
            file_type = self._get_file_type(file_path)
            raw_text = self._extract_text(file_path, file_type)
            
            if not raw_text:
                return None
            
            # Generate summary
            summary = self._generate_summary(raw_text)
            
            # Extract keywords
            keywords = self._extract_keywords(raw_text)
            
            # Count words
            word_count = len(raw_text.split())
            
            return {
                'raw_text': raw_text,
                'summary': summary,
                'keywords': json.dumps(keywords),
                'word_count': word_count
            }
            
        except Exception as e:
            print(f"Error processing document {file_path}: {str(e)}")
            return None
    
    def _get_file_type(self, file_path):
        mime = magic.Magic(mime=True)
        file_type = mime.from_file(file_path)
        
        if 'pdf' in file_type:
            return 'pdf'
        elif 'word' in file_type or file_type == 'application/msword':
            return 'doc'
        elif 'officedocument' in file_type:
            return 'docx'
        elif 'text' in file_type:
            return 'txt'
        else:
            # Fallback to extension
            ext = os.path.splitext(file_path)[1].lower()
            return ext[1:] if ext else 'unknown'
    
    def _extract_text(self, file_path, file_type):
        try:
            if file_type == 'pdf':
                return self._extract_from_pdf(file_path)
            elif file_type == 'doc':
                return self._extract_from_doc(file_path)
            elif file_type == 'docx':
                return self._extract_from_docx(file_path)
            elif file_type == 'txt':
                return self._extract_from_txt(file_path)
            else:
                return None
        except Exception as e:
            print(f"Error extracting text from {file_path}: {str(e)}")
            return None
    
    def _extract_from_pdf(self, file_path):
        text = ""
        with open(file_path, 'rb') as file:
            reader = PdfReader(file)
            for page in reader.pages:
                text += page.extract_text() + "\n"
        return text.strip()
    
    def _extract_from_doc(self, file_path):
        # For .doc files, we'll use antiword if available
        try:
            import subprocess
            result = subprocess.run(['antiword', file_path], capture_output=True, text=True)
            if result.returncode == 0:
                return result.stdout.strip()
        except Exception as e:
            print(f"antiword not available for .doc processing: {str(e)}")
        return None
    
    def _extract_from_docx(self, file_path):
        try:
            doc = docx.Document(file_path)
            text = ""
            for paragraph in doc.paragraphs:
                text += paragraph.text + "\n"
            return text.strip()
        except Exception as e:
            print(f"Error extracting from docx: {str(e)}")
            return None
    
    def _extract_from_txt(self, file_path):
        try:
            with open(file_path, 'r', encoding='utf-8') as file:
                return file.read().strip()
        except UnicodeDecodeError:
            # Try with different encoding
            try:
                with open(file_path, 'r', encoding='latin-1') as file:
                    return file.read().strip()
            except Exception as e:
                print(f"Error reading txt file: {str(e)}")
                return None
    
    def _generate_summary(self, text, num_sentences=3):
        try:
            if not self.nlp:
                return text[:500] + "..." if len(text) > 500 else text
            
            doc = self.nlp(text)
            sentences = [sent.text.strip() for sent in doc.sents]
            
            if len(sentences) <= num_sentences:
                return " ".join(sentences)
            
            # Simple extractive summarization using sentence positions
            # Select first, middle, and last sentences
            summary_sentences = []
            summary_sentences.append(sentences[0])  # First sentence
            
            if len(sentences) > 2:
                middle_idx = len(sentences) // 2
                summary_sentences.append(sentences[middle_idx])  # Middle sentence
            
            summary_sentences.append(sentences[-1])  # Last sentence
            
            return " ".join(summary_sentences[:num_sentences])
            
        except Exception as e:
            print(f"Error generating summary: {str(e)}")
            return text[:500] + "..." if len(text) > 500 else text
    
    def _extract_keywords(self, text, num_keywords=10):
        try:
            if not self.nlp:
                return []
            
            doc = self.nlp(text)
            
            # Extract noun chunks and proper nouns as potential keywords
            keywords = []
            
            # Add named entities
            for ent in doc.ents:
                if len(ent.text.strip()) > 2:
                    keywords.append(ent.text.strip())
            
            # Add noun chunks
            for chunk in doc.noun_chunks:
                if len(chunk.text.strip()) > 2 and not chunk.text.strip().lower() in keywords:
                    keywords.append(chunk.text.strip())
            
            # Remove duplicates and limit
            keywords = list(set(keywords))[:num_keywords]
            
            return keywords
            
        except Exception as e:
            print(f"Error extracting keywords: {str(e)}")
            return []
    
    def calculate_similarity(self, text1, text2):
        try:
            if not text1 or not text2:
                return 0.0
            
            vectorizer = TfidfVectorizer().fit_transform([text1, text2])
            similarity = cosine_similarity(vectorizer[0:1], vectorizer[1:2])[0][0]
            return float(similarity)
            
        except Exception as e:
            print(f"Error calculating similarity: {str(e)}")
            return 0.0
    
    def process_all_documents(self, documents_folder):
        processed_count = 0
        error_count = 0
        
        for filename in os.listdir(documents_folder):
            file_path = os.path.join(documents_folder, filename)
            if os.path.isfile(file_path):
                # Check if already processed
                existing_doc = Document.query.filter_by(filename=filename).first()
                if existing_doc and existing_doc.processed:
                    continue
                
                try:
                    # Process document
                    file_size = os.path.getsize(file_path)
                    file_type = self._get_file_type(file_path)
                    
                    # Create document record
                    if not existing_doc:
                        doc = Document(
                            filename=filename,
                            original_filename=filename,
                            file_path=file_path,
                            file_type=file_type,
                            file_size=file_size,
                            processed=False
                        )
                        db.session.add(doc)
                        db.session.commit()
                    else:
                        doc = existing_doc
                    
                    # Process content
                    content_data = self.process_document(file_path)
                    if content_data:
                        # Update or create content
                        if doc.content:
                            doc.content.raw_text = content_data['raw_text']
                            doc.content.summary = content_data['summary']
                            doc.content.keywords = content_data['keywords']
                            doc.content.word_count = content_data['word_count']
                            doc.content.processing_date = datetime.utcnow()
                        else:
                            content = DocumentContent(
                                document_id=doc.id,
                                **content_data
                            )
                            db.session.add(content)
                        
                        doc.processed = True
                        processed_count += 1
                    else:
                        error_count += 1
                    
                    db.session.commit()
                    
                except Exception as e:
                    print(f"Error processing {filename}: {str(e)}")
                    error_count += 1
                    db.session.rollback()
        
        return processed_count, error_count